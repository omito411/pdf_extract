import os
import re
import pdfplumber
import docx
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = 'supersecretkey'  # Required for flashing messages

UPLOAD_FOLDER = 'uploads/'
DEFAULT_DOCX_TEMPLATE_25 = 'default_template_(2.5mg).docx'
DEFAULT_DOCX_TEMPLATE_5 = 'default_template_(5mg).docx'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Function to extract the NIR Potency Avg (mg) from the PDF
def extract_nir_potency_avg(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()

        # Locate the relevant section based on "Absolute Average"
        nir_potency_avg_start = text.find("Absolute Average")
        if nir_potency_avg_start == -1:
            return None

        # Capture the next 300 characters after the "Absolute Average"
        nir_potency_avg_end = nir_potency_avg_start + 300
        nir_potency_avg_section = text[nir_potency_avg_start:nir_potency_avg_end]

        # Extract the numerical value using regex
        match = re.findall(r"\d+\.\d+", nir_potency_avg_section)
        if match:
            nir_potency_avg_value = match[0]  # First match should be the NIR Potency Avg (mg)
            return float(nir_potency_avg_value)
        else:
            return None

# Function to extract values from the PDF
def extract_values_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        page = pdf.pages[0]
        pdf_text = page.extract_text()

    # Extract values using regex patterns
    nir_qc_pattern = r"QC\d{7,}"
    nir_qc_value = re.search(nir_qc_pattern, pdf_text)
    nir_qc_value = nir_qc_value.group(0) if nir_qc_value else "N/A"

    # Extract the Batch Name (first occurrence in the table)
    batch_name_pattern = r"\b[A-Z]{2,}[0-9]{3,}\b"
    batch_name_match = re.search(batch_name_pattern, pdf_text)
    batch_name = batch_name_match.group(0) if batch_name_match else "N/A"

    lims_mm_pattern = r"LIMS MM Sample\s*(\d+)|:\s*(\d+)"
    lims_mm_sample = re.search(lims_mm_pattern, pdf_text)
    lims_mm_sample = lims_mm_sample.group(0).split()[-1] if lims_mm_sample else "N/A"

    absolute_avg_pattern = r"HPLC Potency Avg \(mg\):.*NIR Potency Avg \(mg\):.*NIR Potency Avg \(%LC\):\s*([0-9]+(?:\.[0-9]*)?)"
    absolute_avg_value = re.search(absolute_avg_pattern, pdf_text)
    absolute_avg_value = absolute_avg_value.group(1) if absolute_avg_value else "N/A"

    specification_pattern = r"Specification <=\s*([0-9]+(?:\.[0-9]*)?)"
    specification_value = re.search(specification_pattern, pdf_text)
    specification_value = specification_value.group(1) if specification_value else "N/A"

    # Extract NIR Potency Avg
    nir_potency_avg_value = extract_nir_potency_avg(pdf_file)

    return {
        'nir_qc': nir_qc_value,
        'batch_name': batch_name,
        'lims_report': lims_mm_sample,
        'absolute_avg': absolute_avg_value,
        'acceptance_criteria': specification_value,
        'nir_potency_avg': nir_potency_avg_value  # Return NIR potency avg
    }

# Function to find NIR potency and validate the file
def validate_potency_value(nir_potency_avg_value, expected_range):
    
    value = nir_potency_avg_value

    if expected_range == '2.5mg' and 2.0 <= value <= 3.0:
        return True
    elif expected_range == '5mg' and 4.5 <= value <= 5.5:
        return True
    return False

# Main function to handle DOCX replacement
def automate_replacement(pdf_file, expected_range, user_qe_number):
    new_values = extract_values_from_pdf(pdf_file)

    if new_values['nir_potency_avg'] is None:
        flash("Error: NIR Potency Avg not found in the PDF.")
        return None

    # Validate the NIR potency based on the expected range
    if not validate_potency_value(new_values['nir_potency_avg'], expected_range):
        flash(f"Incorrect file. Please upload a {expected_range} PDF.")
        return None  # Indicate that the validation failed

    # Determine the appropriate document template
    if expected_range == '2.5mg':
        doc_template = DEFAULT_DOCX_TEMPLATE_25
    else:
        doc_template = DEFAULT_DOCX_TEMPLATE_5

    updated_filename = os.path.join(app.config['UPLOAD_FOLDER'], f"Antaris_verification({expected_range}).docx")

    # Load the DOCX template
    doc = docx.Document(doc_template)

    # Example values to be replaced (assuming user_qe_number)
    old_values = {
        'nir_qc': r"QC\d{7,}",  # Replace QC value
        'batch_name': r"batch:\s*(\w+)",  # Replace batch name
        'lims_report': r"LIMS report\s*\d+|LIMS Report\s*\d+",  # Replace LIMS report value
        'absolute_avg': r"HPLC and NIR potency results was.*?(\d+\.?\d*)%",  # Replace absolute average value
        'acceptance_criteria': r"acceptance criteria.*?≤\s*(\d+\.?\d*)%",  # Replace acceptance criteria value
        'qe_number': r"QE-number"  # To replace with user input
    }

    new_values_mapped = {
        'nir_qc': new_values['nir_qc'],
        'batch_name': f"batch: {new_values['batch_name']}",  # New batch name
        'lims_report': f"LIMS report number {new_values['lims_report']}",
        'absolute_avg': f"absolute average {new_values['absolute_avg']}%",
        'acceptance_criteria': f"acceptance criteria ≤ {new_values['acceptance_criteria']}%",
        'qe_number': user_qe_number if user_qe_number else "QE-number"
    }

    # Replace text in DOCX
    replace_text_in_paragraphs_and_tables(doc, old_values, new_values_mapped)

    # Save the updated DOCX
    doc.save(updated_filename)
    print(f"Document saved as: {updated_filename}")
    return updated_filename

# Function to replace text in paragraphs and tables in the DOCX file
def replace_text_in_paragraphs_and_tables(doc, old_values, new_values):
    for para in doc.paragraphs:
        for key, pattern in old_values.items():
            new_value = new_values[key]
            if re.search(pattern, para.text):
                para.text = re.sub(pattern, new_value, para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, pattern in old_values.items():
                    new_value = new_values[key]
                    if re.search(pattern, cell.text):
                        cell.text = re.sub(pattern, new_value, cell.text)

# Route for file uploads and processing
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        pdf_file = request.files['pdf_file']
        user_qe_number = request.form.get('qe_number')
        potency_type = request.form.get('potency_type')  # '2.5mg' or '5mg'

        if pdf_file and potency_type:
            pdf_filename = secure_filename(pdf_file.filename)
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
            pdf_file.save(pdf_path)

            updated_doc_file = automate_replacement(pdf_path, potency_type, user_qe_number)

            if updated_doc_file:
                return redirect(url_for('download_file', filename=os.path.basename(updated_doc_file)))
            else:
                return redirect(url_for('index'))  # Redirect back to the form if validation failed

    return render_template('index.html')

# Route for downloading the updated DOCX file
@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)


